import os
import sys
import glob
import random
import asyncio
import urllib.request
import docx

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from src.generation.visual_renderer import VisualRenderer
from src.generation.llm_client import UnifiedAIClient
from src.settings import settings

async def run_phase3():
    print("="*50)
    print("🎨 Phase 3: Visual Generation & Layouts (Multimodal + Multi-format)")
    print("="*50)
    
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    samples_dir = os.path.join(base_dir, "output", "samples")
    visuals_dir = os.path.join(samples_dir, "visuals")
    
    md_files = glob.glob(os.path.join(samples_dir, "*.md"))
    if not md_files:
        print("❌ No markdown files found in output/samples/. Run Phase 2 first.")
        return
        
    renderer = VisualRenderer()
    llm_client = UnifiedAIClient()
    await llm_client.initialize()
    
    export_formats = [".png", ".pdf", ".docx", ".txt"]
    
    # Process up to 3 files
    for i, md_path in enumerate(md_files[:3]):
        with open(md_path, "r", encoding="utf-8") as f:
            text = f.read()
            
        doc_type = "Technical Manual" 
        types = list(renderer.profiles.keys())
        for t in types:
            if t.lower() in text.lower():
                doc_type = t
                break
                
        # To show variety, pick a random format for this document
        ext = export_formats[i % len(export_formats)]
        filename_base = os.path.basename(md_path).replace(".md", "")
        
        embedded_img_path = None
        
        # DEMONSTRATION: Generate a drawing/blueprint for the very first document
        if i == 0:
            print(f"🖌️ Generating visual asset for {filename_base} using gpt-image-2...")
            try:
                # Ask Azure DALL-E for a picture
                img_prompt = f"A rough chalk sketch or technical blueprint of an Arcane Industrial entity named {filename_base}. Monochromatic, highly detailed, parchment background."
                response = await llm_client.image_client.images.generate(
                    model=settings.image_generation_deployment,
                    prompt=img_prompt,
                    n=1,
                    size="1024x1024"
                )
                import base64
                img_b64 = response.data[0].b64_json
                
                # Save it locally
                embedded_img_path = os.path.join(visuals_dir, f"{filename_base}_drawing.png")
                with open(embedded_img_path, "wb") as fh:
                    fh.write(base64.b64decode(img_b64))
                print(f"   ✅ Saved generated drawing to {embedded_img_path}")
            except Exception as e:
                print(f"   ❌ Image generation failed: {e}")
        
        out_path = os.path.join(visuals_dir, filename_base + ext)
        print(f"🖼️ Rendering {filename_base} as [{doc_type}] -> {ext}...")
        
        if ext in [".png", ".pdf"]:
            renderer.render_document(text, doc_type, out_path, embedded_image_path=embedded_img_path)
        elif ext == ".docx":
            doc = docx.Document()
            doc.add_heading(filename_base, 0)
            doc.add_paragraph(text)
            if embedded_img_path:
                doc.add_picture(embedded_img_path, width=docx.shared.Inches(5.0))
            doc.save(out_path)
        elif ext == ".txt":
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
                
        print(f"   ✅ Saved to {out_path}")
        
    await llm_client.close()
    print("\n🎉 Phase 3 Sample Generation Complete!")
    print(f"Check {visuals_dir} for the multimodal outputs (including drawing integration).")

if __name__ == "__main__":
    asyncio.run(run_phase3())
