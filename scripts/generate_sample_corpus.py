import os
import sys
import json
import random
import asyncio
import base64
import docx
import time
import logging
import argparse
from datetime import datetime
from tqdm.asyncio import tqdm
import networkx as nx

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from src.graph.config import WorldConfig
from src.graph.generator import KnowledgeGraphGenerator
from src.generation.llm_client import UnifiedAIClient
from src.generation.document_compiler import DocumentCompiler
from src.generation.visual_renderer import VisualRenderer
from src.settings import settings

TEXT_SEMAPHORE = asyncio.Semaphore(10)
IMAGE_SEMAPHORE = asyncio.Semaphore(3)
metrics = {"llm_text": [], "llm_image": [], "render": []}

def setup_logger(corpus_dir):
    log_file = os.path.join(corpus_dir, "generation.log")
    logger = logging.getLogger("SynthloreGenerator")
    logger.setLevel(logging.INFO)
    if logger.hasHandlers():
        logger.handlers.clear()
    fh = logging.FileHandler(log_file)
    fh.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logger.addHandler(fh)
    return logger

async def process_document(llm_client, compiler, renderer, G, node_id, ext, doc_type, corpus_dir, needs_image, logger, theme, world_prompt):
    node_name = G.nodes[node_id]['name']
    
    # Ensure processed directory exists
    processed_dir = os.path.join(corpus_dir, "processed")
    os.makedirs(processed_dir, exist_ok=True)
    
    out_path = os.path.join(processed_dir, f"{node_name}{ext}")
    
    # Resumption check
    if os.path.exists(out_path):
        logger.info(f"[{node_name}] Skipping (already exists).")
        return node_name
        
    # Text Generation
    t0 = time.time()
    async with TEXT_SEMAPHORE:
        for attempt in range(3):
            try:
                text = await compiler.compile_document(G, node_id, doc_type=doc_type, world_prompt=world_prompt)
                break
            except Exception as e:
                if "429" in str(e) or "Too Many Requests" in str(e):
                    logger.warning(f"[{node_name}] 429 Rate Limit on Text. Retrying (attempt {attempt+1})...")
                    await asyncio.sleep(2 ** attempt)
                elif attempt == 2:
                    logger.error(f"[{node_name}] Text failed after retries: {e}")
                    return None
                    
        t_text = time.time() - t0
        metrics["llm_text"].append(t_text)
        
        # Save raw outputs to the main corpus_dir (not processed)
        with open(os.path.join(corpus_dir, f"{node_name}_RAW.md"), "w") as f:
            f.write(text)
        with open(os.path.join(corpus_dir, f"{node_name}_CONTEXT.txt"), "w") as f:
            f.write(compiler.extract_subgraph_context(G, node_id))
            
    # Image Generation
    embedded_img_path = None
    if needs_image:
        t0 = time.time()
        async with IMAGE_SEMAPHORE:
            for attempt in range(3):
                try:
                    node_data = G.nodes[node_id]
                    # Dynamically inject quantitative facts for Track 1A
                    quant_facts = []
                    for k, v in node_data.items():
                        if isinstance(v, (int, float)) and k not in ["established"]:
                            quant_facts.append(f"{k.upper()}: {v}")
                            
                    if quant_facts:
                        fact_str = " | ".join(quant_facts)
                        chart_type = random.choice(["bar chart", "line graph", "radar chart", "scatter plot dashboard"])
                        
                        theme_aesthetics = {
                            "arcane": "sepia-toned, parchment, steampunk, brass and ink",
                            "cyberpunk": "neon, holographic, high-contrast dark mode, glitch art",
                            "space": "utilitarian, monochrome, CRT monitor glow, industrial"
                        }
                        aesthetic = theme_aesthetics.get(theme, "clean corporate")
                        
                        img_prompt = (
                            f"A {aesthetic} {chart_type} representing operational data for '{node_name}'. "
                            f"The following text MUST be explicitly and clearly written inside the chart: '{fact_str}'. "
                            "Make the typography large, highly legible, and centered. The visual style must be deeply immersive to the theme."
                        )
                    else:
                        blueprint_type = random.choice(["technical blueprint", "chalk sketch", "architectural schematic", "holographic cross-section"])
                        theme_aesthetics = {
                            "arcane": "da Vinci style, ink on old vellum, clockwork",
                            "cyberpunk": "wireframe, glowing cyan and magenta, corporate espionage style",
                            "space": "stark white lines on dark blue grid, military grade"
                        }
                        aesthetic = theme_aesthetics.get(theme, "standard engineering")
                        
                        img_prompt = (
                            f"A {aesthetic} {blueprint_type} of a system component named '{node_name}'. "
                            f"Include structural annotations and clear textual labels referencing the entity's name."
                        )
                        
                    response = await llm_client.image_client.images.generate(
                        model=settings.image_generation_deployment,
                        prompt=img_prompt,
                        n=1,
                        size="1024x1024"
                    )
                    img_b64 = response.data[0].b64_json
                    
                    # Save for embedding
                    embedded_img_path = os.path.join(corpus_dir, f"{node_name}_drawing.png")
                    with open(embedded_img_path, "wb") as fh:
                        fh.write(base64.b64decode(img_b64))
                        
                    # Save as a standalone asset in the processed directory for the competitors
                    asset_suffix = "Chart" if quant_facts else "Blueprint"
                    standalone_path = os.path.join(processed_dir, f"{node_name}_{asset_suffix}.png")
                    with open(standalone_path, "wb") as fh:
                        fh.write(base64.b64decode(img_b64))
                        
                    break
                except Exception as e:
                    if "429" in str(e) or "Too Many Requests" in str(e):
                        logger.warning(f"[{node_name}] 429 Rate Limit on Image. Retrying (attempt {attempt+1})...")
                        await asyncio.sleep(4 ** attempt)
                    elif attempt == 2:
                        logger.error(f"[{node_name}] Image failed after retries: {e}")
        
        if embedded_img_path:
            t_img = time.time() - t0
            metrics["llm_image"].append(t_img)
            logger.info(f"[{node_name}] Asset Downloaded: {t_img:.2f}s")
            
    # Rendering
    t0 = time.time()
    
    if ext in [".png", ".pdf"]:
        renderer.render_document(text, doc_type, out_path, embedded_image_path=embedded_img_path)
    elif ext == ".docx":
        doc = docx.Document()
        doc.add_heading(f"{node_name} ({doc_type})", 0)
        doc.add_paragraph(text)
        if embedded_img_path:
            doc.add_picture(embedded_img_path, width=docx.shared.Inches(5.0))
        doc.save(out_path)
    elif ext == ".txt":
        with open(out_path, "w") as f:
            f.write(text)
            
    t_render = time.time() - t0
    metrics["render"].append(t_render)
    logger.info(f"[{node_name}] Rendered ({ext}): {t_render:.2f}s")
    return node_name

async def generate_corpus(num_docs=20, resume_dir=None, theme="arcane", world_prompt=None):
    print("="*60)
    print("🏭 SYNTHLORE PIPELINE: GENERATING FRESH SAMPLE CORPUS")
    print("="*60)
    
    pipeline_start = time.time()
    metrics.clear()
    metrics.update({"llm_text": [], "llm_image": [], "render": []})
    
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Check resumption
    if resume_dir:
        corpus_dir = os.path.abspath(resume_dir)
        if not os.path.exists(corpus_dir):
            raise FileNotFoundError(f"Resume directory not found: {corpus_dir}")
        print(f"🔄 RESUMING PREVIOUS RUN from {corpus_dir}")
        
        # Load Graph
        graph_path = os.path.join(corpus_dir, "ground_truth_graph.json")
        with open(graph_path, "r") as f:
            data = json.load(f)
            G = nx.node_link_graph(data)
            
        # Load Manifest
        manifest_path = os.path.join(corpus_dir, "manifest.json")
        with open(manifest_path, "r") as f:
            manifest = json.load(f)
            
        theme = manifest.get("theme", "arcane")
        world_prompt = manifest.get("world_prompt", world_prompt)
        
        if theme == "cyberpunk":
            config = WorldConfig.default_cyberpunk_corporate()
        elif theme == "space":
            config = WorldConfig.default_deep_space_colony()
        else:
            config = WorldConfig.default_arcane_industrial()
            
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        corpus_dir = os.path.join(base_dir, "output", f"sample_{theme}_{timestamp}")
        os.makedirs(corpus_dir, exist_ok=True)
        
        print(f"\n[Phase 1] Generating graph (Theme: {theme})...")
        if theme == "cyberpunk":
            config = WorldConfig.default_cyberpunk_corporate()
        elif theme == "space":
            config = WorldConfig.default_deep_space_colony()
        else:
            config = WorldConfig.default_arcane_industrial()
            
        generator = KnowledgeGraphGenerator(config)
        G = generator.generate(num_nodes=num_docs + 10)
        
        graph_path = os.path.join(corpus_dir, "ground_truth_graph.json")
        with open(graph_path, "w") as f:
            json.dump(nx.node_link_data(G), f, indent=2)
            
        # Create Manifest for deterministic resumption
        dist = config.corpus_distribution
        num_with_image = int(num_docs * dist.image_injection_ratio)
        formats = []
        for ext, ratio in dist.format_ratios.items():
            formats.extend([ext] * int(num_docs * ratio))
        while len(formats) < num_docs:
            formats.append(list(dist.format_ratios.keys())[0])
        random.shuffle(formats)
        
        target_nodes = list(G.nodes)[:num_docs]
        valid_image_indices = [i for i, ext in enumerate(formats) if ext != ".txt"]
        image_indices = set(random.sample(valid_image_indices, min(num_with_image, len(valid_image_indices))))
        
        def generate_dynamic_doc_type(node_data, theme):
            bases = ["Log", "Report", "Ledger", "Manifest", "Journal", "Diary", "Contract", "Directive", "Notice", "Intercept", "Cache", "Record", "Transcript", "Missive", "Audit", "Blueprint Overview", "Maintenance Schedule", "Grievance", "Requisition Form", "Interrogation Transcript", "Prophecy", "Heretical Pamphlet", "Manifesto", "Smuggled Cipher"]
            faction = node_data.get("faction")
            role = node_data.get("role") or node_data.get("manager") or node_data.get("specialty")
            node_type = node_data.get("type", "")
            prefixes = []
            if faction and faction != "UNKNOWN_FACTION": prefixes.append(faction)
            if role and not role.startswith("Unknown"): prefixes.append(f"{role}'s")
            if node_type and not node_type.startswith("Unknown") and node_type not in ["Person", "Operative", "Colonist", "Commander", "Executive", "Overseer"]:
                prefixes.append(node_type)
            prefix = random.choice(prefixes) if prefixes else random.choice(["Classified", "Standard", "Priority", "Encrypted"])
            return f"{prefix} {random.choice(bases)}"
        
        manifest = {"theme": theme, "world_prompt": world_prompt} # Store in manifest for resumption
        for i, node_id in enumerate(target_nodes):
            manifest[node_id] = {
                "ext": formats[i],
                "doc_type": generate_dynamic_doc_type(G.nodes[node_id], theme),
                "needs_image": (i in image_indices)
            }
            
        with open(os.path.join(corpus_dir, "manifest.json"), "w") as f:
            json.dump(manifest, f, indent=2)
            
    # Setup LLM and Compiler
    llm_client = UnifiedAIClient()
    await llm_client.initialize()
    compiler = DocumentCompiler(llm_client, config)
    
    # Setup Renderer
    renderer = VisualRenderer()
    
    logger = setup_logger(corpus_dir)
    print(f"\n[Phase 2 & 3] Compiling and Rendering {len(manifest) - 2} Documents Concurrently...")
    
    tasks = []
    for node_id, m_data in manifest.items():
        if node_id in ["theme", "world_prompt"]:
            continue
        tasks.append(asyncio.create_task(
            process_document(
                llm_client, compiler, renderer, G, node_id, 
                m_data["ext"], m_data["doc_type"], corpus_dir, m_data["needs_image"], logger, theme, world_prompt
            )
        ))
        
    for coro in tqdm.as_completed(tasks, total=len(tasks), desc="Processing Documents"):
        await coro

    # Reporting
    total_time = time.time() - pipeline_start
    print("\n" + "="*60)
    print("📊 PARALLEL PIPELINE METRICS")
    print("="*60)
    print(f"Output Directory: {corpus_dir}")
    print(f"Execution Time (This Run): {total_time:.2f}s")
    if metrics['llm_text']:
        print(f"Average Text Gen:  {sum(metrics['llm_text'])/len(metrics['llm_text']):.2f}s")
    if metrics['llm_image']:
        print(f"Average Image Gen: {sum(metrics['llm_image'])/len(metrics['llm_image']):.2f}s")
    if metrics['render']:
        print(f"Average Rendering: {sum(metrics['render'])/len(metrics['render']):.2f}s")
    print("="*60)
    print("🎉 CORPUS COMPLETE!\n")
    await llm_client.close()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate synthetic corporate lore corpus.")
    parser.add_argument("--num_docs", type=int, default=20, help="Number of documents to generate")
    parser.add_argument("--resume", type=str, default=None, help="Directory path of a previous run to resume")
    parser.add_argument("--theme", type=str, choices=["arcane", "cyberpunk", "space"], default="arcane", help="Theme setting for the corpus")
    parser.add_argument("--world_prompt", type=str, default=None, help="Custom detailed prompt to enforce specific world building elements")
    args = parser.parse_args()
    
    asyncio.run(generate_corpus(num_docs=args.num_docs, resume_dir=args.resume, theme=args.theme, world_prompt=args.world_prompt))
